import os
import pypandoc
import markdown
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import html2text
from html.parser import HTMLParser
from io import StringIO

def markdown_to_rtf(markdown_text):
    """Convert markdown text to RTF format with enhanced styling"""
    try:
        # First, normalize line endings and ensure proper spacing between sections
        markdown_text = markdown_text.replace('\r\n', '\n')
        
        # Make sure headers have proper spacing before them
        for header_level in range(1, 7):
            header_marker = '#' * header_level + ' '
            markdown_text = re.sub(r'([^\n])\n' + re.escape(header_marker), 
                                  r'\1\n\n' + header_marker, 
                                  markdown_text)
        
        # Convert markdown to HTML with essential extensions
        html = markdown.markdown(markdown_text, extensions=['extra', 'tables'])
        
        # Define RTF header with optimized styling
        rtf_header = [
            r"{\rtf1\ansi\ansicpg1252\cocoartf2761",
            r"\cocoatextscaling0\cocoaplatform0",
            r"{\fonttbl",
            r"\f0\fswiss\fcharset0 Helvetica;",
            r"\f1\fswiss\fcharset0 Helvetica-Bold;",
            r"\f2\fswiss\fcharset0 Helvetica-Italic;",
            r"\f3\fswiss\fcharset0 Helvetica-BoldItalic;",
            r"}",
            r"{\colortbl;\red0\green0\blue0;\red0\green0\blue255;}",
            r"\paperw12240\paperh15840",
            r"\margl1440\margr1440",
            r"\vieww12000\viewh15000\viewkind0",
            r"\pard\tx720\pardeftab720\partightenfactor0",
            r"\f0\fs24",
        ]
        
        rtf = '\n'.join(rtf_header)
        
        class RTFConverter(HTMLParser):
            def __init__(self):
                super().__init__()
                self.result = StringIO()
                self.in_paragraph = False
                self.in_header = False
                self.header_level = 0
                self.in_list = False
                self.in_list_item = False
                self.in_bold = False
                self.in_italic = False
                self.skip_next_data = False
                self.in_link = False
                self.link_content = ""
                self.tags_stack = []
                
            def handle_starttag(self, tag, attrs):
                self.tags_stack.append(tag)
                
                if tag == 'h1':
                    self.in_header = True
                    self.header_level = 1
                    self.result.write('\\par\\pard\\f1\\fs40\\b ')
                elif tag == 'h2':
                    self.in_header = True
                    self.header_level = 2
                    self.result.write('\\par\\pard\\f1\\fs32\\b ')
                elif tag == 'h3':
                    self.in_header = True
                    self.header_level = 3
                    self.result.write('\\par\\pard\\f1\\fs28\\b ')
                elif tag == 'h4':
                    self.in_header = True
                    self.header_level = 4
                    self.result.write('\\par\\pard\\f1\\fs24\\b ')
                elif tag == 'p':
                    if not self.in_paragraph and not self.in_list_item:
                        self.result.write('\\par\\pard ')
                    self.in_paragraph = True
                elif tag == 'ul' or tag == 'ol':
                    self.in_list = True
                    if not self.in_paragraph:
                        self.result.write('\\par\\pard ')
                elif tag == 'li':
                    self.in_list_item = True
                    self.result.write('\\par\\pard\\fi-360\\li720 {\\bullet} ')
                elif tag == 'strong' or tag == 'b':
                    self.in_bold = True
                    self.result.write('\\b ')
                elif tag == 'em' or tag == 'i':
                    self.in_italic = True
                    self.result.write('\\i ')
                elif tag == 'a':
                    self.in_link = True
                    href = next((attr[1] for attr in attrs if attr[0] == 'href'), '')
                    if href:
                        # Format as blue underlined text
                        self.result.write('\\cf2\\ul ')
                elif tag == 'br':
                    self.result.write('\\par ')
                elif tag == 'table':
                    self.result.write('\\par\\pard ')
                elif tag == 'tr':
                    self.result.write('\\par\\pard ')
                elif tag == 'td' or tag == 'th':
                    self.result.write('\\cell ')
                
            def handle_endtag(self, tag):
                if self.tags_stack and self.tags_stack[-1] == tag:
                    self.tags_stack.pop()
                
                if tag in ('h1', 'h2', 'h3', 'h4'):
                    self.in_header = False
                    self.header_level = 0
                    self.result.write('\\f0\\b0\\fs24\\par\\par\\pard ')
                elif tag == 'p':
                    self.in_paragraph = False
                    if not self.in_list_item:
                        self.result.write('\\par\\par\\pard ')
                elif tag in ('ul', 'ol'):
                    self.in_list = False
                    self.result.write('\\par\\pard ')
                elif tag == 'li':
                    self.in_list_item = False
                    self.result.write('\\par ')
                elif tag == 'strong' or tag == 'b':
                    self.in_bold = False
                    self.result.write('\\b0 ')
                elif tag == 'em' or tag == 'i':
                    self.in_italic = False
                    self.result.write('\\i0 ')
                elif tag == 'a':
                    self.in_link = False
                    self.result.write('\\cf1\\ulnone ')
                elif tag == 'tr':
                    self.result.write('\\row ')
                elif tag == 'table':
                    self.result.write('\\par\\pard ')
                
            def handle_data(self, data):
                if self.skip_next_data:
                    self.skip_next_data = False
                    return
                
                # Handle RTF escaping for special characters
                data = data.replace('\\', '\\\\')
                data = data.replace('{', '\\{')
                data = data.replace('}', '\\}')
                
                # Handle Unicode characters
                replacements = {
                    '•': '{\\bullet}',
                    '"': '\\"',
                    '"': '\\"',
                    ''': "\\'92",
                    ''': "\\'92",
                    '…': '...',
                    '–': '-',
                    '—': '--',
                    '\u2022': '{\\bullet}',  # Unicode bullet
                    '\u2018': "\\'91",       # Left single quote
                    '\u2019': "\\'92",       # Right single quote
                    '\u201C': "\\'93",       # Left double quote
                    '\u201D': "\\'94",       # Right double quote
                }
                
                for old, new in replacements.items():
                    data = data.replace(old, new)
                
                self.result.write(data)
            
            def get_rtf(self):
                return self.result.getvalue()
        
        # Parse HTML to RTF
        converter = RTFConverter()
        converter.feed(html)
        rtf_body = converter.get_rtf()
        
        # Add proper paragraph spacing
        rtf_body = rtf_body.replace('\\par\\par', '\\par\\sa180\\par')
        
        # Special handling for product links with bold-italic formatting [***text***]
        rtf_body = re.sub(r'\\\*\\\*\\\*([^*]+)\\\*\\\*\\\*', r'\\b\\i \1\\i0\\b0 ', rtf_body)
        
        # Add content and close RTF
        rtf += rtf_body + "\n}"
        
        return rtf
        
    except Exception as e:
        print(f"Error in RTF conversion: {str(e)}")
        return str(e)

def enhance_rtf_formatting(rtf_content):
    """Enhance the RTF formatting with custom styling"""
    # Add custom font table
    font_table = (
        r"{\fonttbl"
        r"{\f0\fswiss\fcharset0 Helvetica-Bold;}"
        r"{\f1\fswiss\fcharset0 Helvetica;}"
        r"{\f2\fmodern\fcharset0 Menlo-Regular;}"
        r"}"
    )
    
    # Replace default font table
    rtf_content = re.sub(r'{\fonttbl.*?}', font_table, rtf_content)
    
    # Enhance heading styles
    rtf_content = re.sub(r'\\f0\\fs48', r'\\f0\\fs40\\b', rtf_content)  # H1
    rtf_content = re.sub(r'\\f0\\fs36', r'\\f0\\fs32\\b', rtf_content)  # H2
    rtf_content = re.sub(r'\\f0\\fs28', r'\\f0\\fs28\\b', rtf_content)  # H3
    
    return rtf_content

def basic_markdown_to_rtf(markdown_text):
    """Fallback basic markdown to RTF converter"""
    # First convert markdown to HTML
    html = markdown.markdown(markdown_text)
    
    # Basic RTF header with more complete styling
    rtf = [
        r"{\rtf1\ansi\ansicpg1252\cocoartf2639\cocoatextscaling0",
        r"{\fonttbl\f0\fswiss\fcharset0 Helvetica;\f1\froman\fcharset0 TimesNewRoman;}",
        r"{\colortbl;\red0\green0\blue0;\red0\green0\blue255;}",
        r"\paperw12240\paperh15840\margl1440\margr1440\vieww12000\viewh15000\viewkind0",
        r"\pard\tx720\tx1440\tx2160\tx2880\tx3600\tx4320\tx5040\tx5760\tx6480\tx7200\tx7920\tx8640\pardirnatural\partightenfactor0",
        r"\f0\fs24\fsmilli12200"
    ]
    rtf = '\n'.join(rtf)
    
    # Convert HTML to RTF
    rtf_body = html
    
    # Headers
    rtf_body = re.sub(r'<h1>(.*?)</h1>', r'\\par\\f1\\fs36\\b \1\\b0\\fs24\\f0\\par\n', rtf_body)
    rtf_body = re.sub(r'<h2>(.*?)</h2>', r'\\par\\f1\\fs32\\b \1\\b0\\fs24\\f0\\par\n', rtf_body)
    rtf_body = re.sub(r'<h3>(.*?)</h3>', r'\\par\\f1\\fs28\\b \1\\b0\\fs24\\f0\\par\n', rtf_body)
    
    # Paragraphs
    rtf_body = re.sub(r'<p>(.*?)</p>', r'\\par \1\\par\n', rtf_body)
    
    # Lists
    rtf_body = re.sub(r'<ul>(.*?)</ul>', r'\\par \1\\par\n', rtf_body)
    rtf_body = re.sub(r'<li>(.*?)</li>', r'\\par • \1', rtf_body)
    
    # Inline formatting
    rtf_body = re.sub(r'<strong>(.*?)</strong>', r'\\b \1\\b0 ', rtf_body)
    rtf_body = re.sub(r'<em>(.*?)</em>', r'\\i \1\\i0 ', rtf_body)
    rtf_body = re.sub(r'<code>(.*?)</code>', r'\\f1 \1\\f0 ', rtf_body)
    
    # Links
    rtf_body = re.sub(r'<a href="(.*?)">(.*?)</a>', r'\\cf2\\ul \2\\cf1\\ulnone', rtf_body)
    
    # Clean up any remaining HTML tags
    rtf_body = re.sub(r'<[^>]+>', '', rtf_body)
    
    # Escape special characters
    rtf_body = rtf_body.replace('\\', '\\\\')
    rtf_body = rtf_body.replace('{', '\\{')
    rtf_body = rtf_body.replace('}', '\\}')
    
    # Fix double escaping of RTF commands
    rtf_body = rtf_body.replace('\\\\par', '\\par')
    rtf_body = rtf_body.replace('\\\\b', '\\b')
    rtf_body = rtf_body.replace('\\\\i', '\\i')
    rtf_body = rtf_body.replace('\\\\f', '\\f')
    rtf_body = rtf_body.replace('\\\\fs', '\\fs')
    
    # Add converted body and close RTF
    rtf += rtf_body + "\n}"
    
    return rtf

def markdown_to_docx(markdown_text):
    """Convert markdown text to DOCX format"""
    # Convert markdown to HTML first
    html = markdown.markdown(markdown_text)
    
    # Create a new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Helvetica'
    style.font.size = Pt(11)
    
    # Convert HTML to plain text while preserving basic formatting
    h = html2text.HTML2Text()
    h.body_width = 0  # Disable wrapping
    text = h.handle(html)
    
    # Process the text line by line
    current_list = None
    for line in text.split('\n'):
        if line.startswith('# '):  # H1
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.font.size = Pt(20)
            run.font.bold = True
            p.space_after = Pt(12)
        elif line.startswith('## '):  # H2
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.font.size = Pt(16)
            run.font.bold = True
            p.space_after = Pt(10)
        elif line.startswith('### '):  # H3
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.font.size = Pt(14)
            run.font.bold = True
            p.space_after = Pt(8)
        elif line.startswith('* '):  # List items
            if current_list is None:
                current_list = doc.add_paragraph()
            p = current_list.add_run('\n• ' + line[2:])
        elif line.strip() == '':  # Empty line
            current_list = None
            doc.add_paragraph()
        else:  # Regular paragraph
            current_list = None
            p = doc.add_paragraph(line)
            p.space_after = Pt(10)
    
    return doc

def save_as_docx(doc, filename):
    """Save the document as DOCX"""
    # Ensure the output directory exists
    os.makedirs(os.path.dirname(filename), exist_ok=True)
    doc.save(filename)

# Example usage
def convert_markdown_file(input_file, output_file_rtf, output_file_docx=None):
    """Convert a markdown file to both RTF and optionally DOCX formats"""
    try:
        # Read the markdown file
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Preprocess markdown to ensure proper spacing
        markdown_content = markdown_content.replace('\r\n', '\n')
        
        # Convert to RTF
        rtf_content = markdown_to_rtf(markdown_content)
        
        # Save RTF file
        with open(output_file_rtf, 'w', encoding='utf-8') as f:
            f.write(rtf_content)
        
        # Optionally convert to DOCX
        if output_file_docx:
            doc = markdown_to_docx(markdown_content)
            save_as_docx(doc, output_file_docx)
        
        return True
    except Exception as e:
        print(f"Error converting file: {str(e)}")
        return False

# Example:
# convert_markdown_file('input.md', 'output.rtf', 'output.docx')